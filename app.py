import streamlit as st
import pdfplumber
import sys
import os
from io import BytesIO
from docx import Document
from gtts import gTTS
from streamlit_mic_recorder import mic_recorder

# --- 1. PATH RESILIENCE ---
root_path = os.path.dirname(os.path.abspath(__file__))
if root_path not in sys.path:
    sys.path.insert(0, root_path)

# --- 2. PAGE CONFIGURATION ---
st.set_page_config(page_title="Career Orchestrator v2.0", layout="wide")

# --- 3. IMPORT GENERALIZED AGENTS ---
try:
    from agents.strategy_arch import run_strategy_architect
    from agents.ats_architect import run_ats_architect
    from agents.human_narrator import run_human_narrator
    
    # FIX: Import both the Guardian AND the Evaluator from the integrity agent
    # this allows your 99.9% USPS accuracy stories to be graded [cite: 2026-03-23]
    from agents.integrity import run_integrity_guardian, evaluate_and_reorg_answer 
    
except Exception as e:
    st.error(f"🛑 Critical Load Failure: {e}")
    st.stop()

# --- 4. SESSION STATE ---
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'final_results' not in st.session_state:
    st.session_state.final_results = None

# --- 5. HELPERS: DOWNLOAD ENGINE ---
def generate_docx_report(company, level, jd, results):
    doc = Document()
    doc.add_heading(f"Strategy Report: {company}", 0)
    
    # Layer 1: Strategy
    strat = results.get('strategy', {})
    doc.add_heading("Layer 1: Strategic Audit", level=1)
    doc.add_paragraph(f"Match Score: {strat.get('match_score', 'N/A')}%")
    doc.add_paragraph(strat.get('learning_syllabus', ''))
    
    # Layer 2: ATS
    ats = results.get('ats', {})
    doc.add_heading("Layer 2: Optimized ATS Bullets", level=1)
    for exp in ats.get('optimized_bullets', []):
        doc.add_heading(exp.get('Role', 'Experience'), level=2)
        for bullet in exp.get('Bullets', []):
            doc.add_paragraph(bullet, style='List Bullet')
    
    # Layer 3: Formal Cover Letter (Restored V1 Style)
    doc.add_heading("Layer 3: Formal Cover Letter", level=1)
    doc.add_paragraph(results.get('narrative', {}).get('cover_letter_narrative', ''))
    
    # Layer 4: Interview Prep
    doc.add_heading("Layer 4: Interview Drills", level=1)
    questions = results.get('integrity', {}).get('interview_questions', {})
    for q_type, q_text in questions.items():
        doc.add_heading(q_type.replace('_', ' ').title(), level=2)
        doc.add_paragraph(q_text)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 6. SIDEBAR: PERSONA & DNA ---
with st.sidebar:
    st.header("👤 Persona Discovery")
    api_key = st.text_input("DeepSeek API Key", type="password")
    user_strengths = st.text_area("Your Top 3 Strengths", placeholder="e.g., Reliability, Python, Grit")
    user_weaknesses = st.text_area("Gaps/Weaknesses", placeholder="e.g., Impatience")
    writing_dna_choice = st.selectbox("Base Writing Style", ["Blunt & Gritty", "Professional", "Academic"])
    
    st.header("🧬 Linguistic DNA (V1 Style)")
    style_file = st.file_uploader("Upload a Writing Sample (PDF/Docx)", type=["pdf", "docx"], key="style_upload")
    style_text = ""
    if style_file:
        try:
            if style_file.type == "application/pdf":
                with pdfplumber.open(style_file) as pdf:
                    style_text = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
            else:
                doc_style = Document(style_file)
                style_text = "\n".join([p.text for p in doc_style.paragraphs])
            st.success("✅ DNA Captured")
        except: st.error("Style parsing failed.")

    st.header("📂 Ingestion")
    uploaded_file = st.file_uploader("Upload Master Resume", type=["pdf", "docx"])
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                st.session_state.resume_text = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
        else:
            doc_res = Document(uploaded_file)
            st.session_state.resume_text = "\n".join([p.text for p in doc_res.paragraphs])
        st.success("✅ Resume Memory Locked")

# --- 7. MAIN UI ---
st.title("🚀 Career Orchestrator: Multi-Agent Platform")
c1, c2 = st.columns(2)
with c1:
    company_name = st.text_input("Target Company", value="Quva Pharma")
    job_level = st.selectbox("Job Level", ["Intern", "Junior", "Senior", "Lead"]) # Dynamic Variable
with c2:
    jd_input = st.text_area("Target Job Description", height=150)

# --- 8. ORCHESTRATION ---
if st.button("🔥 Run Full Optimization"):
    if not st.session_state.resume_text or not api_key:
        st.error("❌ Missing Resume or API Key.")
    else:
        try:
            with st.status("Orchestrating Agents...") as status:
                st.write("🕵️ Phase 1: Strategy Architect...")
                strat = run_strategy_architect(st.session_state.resume_text, jd_input, job_level, company_name, api_key, user_strengths, user_weaknesses, writing_dna_choice)
                
                st.write("🤖 Phase 2: ATS Architect...")
                ats = run_ats_architect(st.session_state.resume_text, strat.get('missing_gaps', []), jd_input, strat.get('persona_assessment', ''), job_level, company_name, api_key)
                
                st.write("✍️ Phase 3: Human Narrator...")
                # Dynamic call with style DNA
                narrative = run_human_narrator(st.session_state.resume_text, jd_input, strat.get('persona_assessment', ''), writing_dna_choice, company_name, job_level, api_key, style_text)
                
                st.write("🛡️ Phase 4: Integrity Guardian...")
                integrity = run_integrity_guardian(st.session_state.resume_text, ats, narrative, strat.get('missing_gaps', []), api_key)
                
                st.session_state.final_results = {"strategy": strat, "ats": ats, "narrative": narrative, "integrity": integrity}
                status.update(label="✅ Optimization Complete!", state="complete")
            st.rerun() 
        except Exception as e:
            st.error(f"❌ Orchestration Error: {e}")
            
# --- 9. RESULTS DISPLAY ---
if st.session_state.final_results:
    res = st.session_state.final_results
    st.divider()
    
    # 🎯 PERSONA & INTERNAL STRATEGY (V1 Restoration)
    st.subheader("👤 Candidate Persona & Management Roadmap")
    col_narr, col_risk = st.columns(2)
    with col_narr:
        st.info(f"**Persona Fit:** {res['strategy'].get('persona_assessment', 'Standard')}")
    with col_risk:
        # Fixed alignment
        st.warning(res['narrative'].get('internal_placement_strategy', "Internal Strategy processing..."))

    t1, t2, t3, t4 = st.tabs(["📊 Strategy", "📄 ATS Bullets", "✉️ Formal Cover Letter", "🎙️ Voice Practice"])
    
    with t1:
        st.metric("Match Score", f"{res['strategy'].get('match_score', 0)}%")
        st.markdown(res['strategy'].get('learning_syllabus', ''))
    
    with t2:
        st.subheader("🎯 ATS-Optimized Impact Bullets")
        
        # Get the 'ats' dictionary from the results
        ats_data = res.get('ats', {})
        
        if ats_data:
            # Match these keys EXACTLY to the agent's prompt
            verdict = ats_data.get('recruiter_scan_verdict', "No verdict generated.")
            keywords = ats_data.get('ats_keywords_hit', [])
            
            st.success(f"**Recruiter Scan Verdict:** {verdict}")
            st.write(f"**Keywords Infiltrated:** {', '.join(keywords) if keywords else 'None identified.'}")
            
            st.divider()
    
            experience_list = ats_data.get('optimized_experience', [])
            if experience_list:
                for item in experience_list:
                    # Logic: Use 'Role' or 'Job Title' fallback
                    role_title = item.get('Role') or item.get('Job Title') or "Experience Entry"
                    with st.expander(f"📂 {role_title}"):
                        for bullet in item.get('Bullets', []):
                            st.write(bullet)
            else:
                st.warning("No experience bullets were generated.")
        else:
            st.error("ATS Data missing.")

    with t3:
        st.subheader("✉️ Formal Persona-Driven Cover Letter")
        letter = res['narrative'].get('cover_letter_narrative', "")
        st.text_area("Review your Letter:", letter, height=400)
        st.download_button("📥 Download Report (.docx)", generate_docx_report(company_name, job_level, jd_input, res), file_name=f"{company_name}_Career_Pack.docx")

    with t4:
        st.subheader("🎙️ Interactive Technical Drill")
        # Fixed internal indent here
        questions = res['integrity'].get('interview_questions', {})
        
        if questions:
            q_selected = st.selectbox("Select Drill:", list(questions.values()))
            st.info(f"**Challenge:** {q_selected}")
            
            with st.expander("💡 View Strategic Hint (STAR Method)"):
                st.write("To nail this for the A3 Team, focus on:")
                st.write("- **Situation:** Briefly describe the 70,000-row survey or AGENT.AI context [cite: 2026-01-09, 2026-03-11].")
                st.write("- **Task:** What was the specific data bottleneck (e.g., messy SQL joins)?")
                st.write("- **Action:** Use 'Workhorse' verbs: *Engineered, Normalized, Architected*.")
                st.write("- **Result:** Mention the 1st Place win or the 99.9% USPS accuracy [cite: 2026-03-23].")
            
            if st.button("📢 Hear Question"):
                tts = gTTS(text=q_selected, lang='en')
                audio_fp = BytesIO()
                tts.write_to_fp(audio_fp)
                st.audio(audio_fp.getvalue(), format='audio/mp3')
            
            st.divider()
            
            st.write("Record your answer:")
            audio_data = mic_recorder(start_prompt="🎤 Start Recording", stop_prompt="🛑 Stop", key='browser_mic')
            
            if audio_data:
                st.audio(audio_data['bytes'])
                if st.button("⚖️ Get Blunt Feedback & STAR Reorg"):
                    with st.spinner("Analyzing your grit..."):
                        # Calling consolidated coach
                        feedback = evaluate_and_reorg_answer(q_selected, "Audio response captured.", api_key)
                        st.markdown("### 📝 Coach's Blunt Feedback")
                        st.warning(feedback)
                        st.success("✅ Review the 'Perfect STAR Answer' above to calibrate your next attempt.")

    
    
